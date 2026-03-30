"""Document processing service for extracting text from .doc and .docx files"""
import logging
import os
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


class DocProcessor:
    """Handles text extraction from .doc and .docx files using python-docx"""

    def extract_text_from_doc(self, file_path: str) -> Optional[str]:
        """
        Extract text from a .doc or .docx file.

        For .docx files, python-docx is used to read paragraphs and tables.
        For legacy .doc files, a best-effort binary read is attempted.

        Args:
            file_path: Path to the .doc/.docx file

        Returns:
            Extracted text as a string, or None if extraction fails
        """
        ext = Path(file_path).suffix.lower()
        try:
            if ext == ".docx":
                return self._extract_docx(file_path)
            elif ext == ".doc":
                return self._extract_doc(file_path)
            else:
                logger.error(f"Unsupported file extension: {ext}")
                return None
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None

    def _extract_docx(self, file_path: str) -> Optional[str]:
        """Extract text from a .docx file using python-docx."""
        try:
            import docx  # python-docx

            doc = docx.Document(file_path)
            parts: list[str] = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        parts.append(" | ".join(row_cells))

            return "\n".join(parts) if parts else None
        except ImportError:
            logger.error(
                "python-docx is not installed. Run: pip install python-docx"
            )
            return None
        except Exception as e:
            logger.error(f"Failed to extract .docx text from {file_path}: {str(e)}")
            return None

    def _extract_doc(self, file_path: str) -> Optional[str]:
        """
        Best-effort text extraction from a legacy .doc (OLE compound) file.

        Uses python-docx2txt when available; falls back to raw byte scan.
        """
        try:
            import docx2txt

            text = docx2txt.process(file_path)
            return text.strip() if text and text.strip() else None
        except ImportError:
            pass  # fall through to raw extraction
        except Exception as e:
            logger.warning(f"docx2txt extraction failed for {file_path}: {str(e)}")

        # Raw fallback: extract printable ASCII runs from the binary
        try:
            with open(file_path, "rb") as f:
                raw = f.read()
            # Filter printable ASCII (32-126) and replace non-printable with space
            text = "".join(chr(b) if 32 <= b <= 126 else " " for b in raw)
            # Collapse whitespace runs
            import re
            text = re.sub(r" {3,}", "  ", text).strip()
            return text if len(text) > 50 else None
        except Exception as e:
            logger.error(f"Raw .doc extraction failed for {file_path}: {str(e)}")
            return None
