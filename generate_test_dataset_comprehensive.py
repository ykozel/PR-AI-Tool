"""
Generate comprehensive test dataset for LLM career development analysis.
Creates multi-year test documents for 10+ distinct employee scenarios:

1. Clear Progression: Junior â†’ Mid â†’ Senior (3-year arc)
2. Regression Then Recovery: High performer decline + comeback
3. Mixed Performance: Strong technical skills, weak leadership
4. Stagnation: No growth, no decline (career plateau)
5. Rapid Growth: Entry-level to advanced in 2 years
6. Lateral Move: Skill transfer with learning curve
7. Specialist Path: Deep expertise in narrowing domain
8. Burnout & Turnaround: Performance crisis and recovery
9. High Performer Plateau: Excellent but not advancing
10. Career Retooling: Technical to management track
11. Part-time/Return: Re-entry into workforce
12. Project Crisis Management: Handling major setback

Run:
    pip install python-docx
    python generate_test_dataset_comprehensive.py
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create output directory
os.makedirs("test_docs_comprehensive", exist_ok=True)


def add_doc_header(doc, doc_type, scenario_name):
    """Add document type and scenario header"""
    header = doc.add_heading(f"[{doc_type}] {scenario_name}", level=0)
    header.runs[0].font.color.rgb = RGBColor(0x70, 0x70, 0x70)  # Gray
    header.runs[0].font.size = Pt(10)
    doc.add_paragraph()  # Spacing


def add_section_type(doc, section_type):
    """Add section type indicator for easy LLM parsing"""
    p = doc.add_paragraph()
    run = p.add_run(f"[{section_type}]")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)  # Medium gray


def add_heading(doc, text, level=1):
    """Add formatted heading to document"""
    h = doc.add_heading(text, level=level)
    if level == 1:
        h.runs[0].font.color.rgb = RGBColor(0x1F, 0x47, 0x88)  # Professional blue
    return h


def add_feedback(doc, source, content):
    """Add feedback section with formatting"""
    p = doc.add_paragraph()
    run = p.add_run(f"{source}: ")
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(content)


def add_list_item(doc, text, indent=0):
    """Add bullet point"""
    p = doc.add_paragraph(text, style='List Bullet')
    return p


# ===== SCENARIO 13: ELENA RODRIGUEZ â€” QA AUTOMATION ENGINEER =====
# Three separate files matching real upload types: CLIENT FEEDBACK, PROJECT FEEDBACK, PDP

def create_elena_client_feedback_2025():
    """Elena Rodriguez 2025 â€” Client Feedback (separate document)"""
    doc = Document()
    add_doc_header(doc, "CLIENT FEEDBACK", "QA Automation Engineer â€” FinCore Banking Platform")
    add_heading(doc, "Client Feedback Form", 1)
    add_heading(doc, "Review Period: January 2025 â€“ December 2025", 2)

    add_heading(doc, "Engagement Details", 2)
    doc.add_paragraph("Employee: Elena Rodriguez")
    doc.add_paragraph("Role: QA Automation Engineer")
    doc.add_paragraph("Project / Engagement: FinCore Banking Platform â€” Automated Testing Stream")
    doc.add_paragraph("Client Organisation: FinCore Capital Ltd")
    doc.add_paragraph("Client Contact: Mark Ellison, VP of Engineering, FinCore Capital Ltd")

    add_section_type(doc, "CLIENT FEEDBACK")
    add_heading(doc, "Satisfaction Ratings", 2)
    add_list_item(doc, "Overall satisfaction: 5 / 5")
    add_list_item(doc, "Quality of deliverables: 5 / 5")
    add_list_item(doc, "Technical competence: 5 / 5")
    add_list_item(doc, "Communication and responsiveness: 4 / 5")
    add_list_item(doc, "Meeting commitments and deadlines: 5 / 5")
    add_list_item(doc, "Proactivity and initiative: 5 / 5")
    add_list_item(doc, "Would you engage this person again: Yes â€” strongly recommended")

    add_heading(doc, "Quality of Work and Technical Competence", 2)
    add_feedback(doc, "Mark Ellison, VP Engineering â€” FinCore Capital",
        "Elena's automation framework has been the single biggest quality improvement on this programme. "
        "She did not just execute what we asked â€” she identified coverage gaps we did not know existed, "
        "proposed a risk-based testing approach for high-value transaction flows, and built a suite that "
        "our internal team can independently maintain. That level of ownership is rare.")

    add_heading(doc, "Communication and Responsiveness", 2)
    add_feedback(doc, "Mark Ellison, VP Engineering â€” FinCore Capital",
        "Communication is good overall. Elena is clear and precise when she explains a defect or risk. "
        "One area to improve: she could surface blockers earlier rather than attempting to resolve them "
        "herself first. When she does escalate, her analysis is always thorough and well-prepared.")

    add_heading(doc, "Meeting Commitments", 2)
    add_feedback(doc, "Mark Ellison, VP Engineering â€” FinCore Capital",
        "All milestones delivered on schedule. The regression suite now catches approximately 95% of "
        "defects before UAT â€” that has fundamentally changed how confident we feel going into each release.")

    add_heading(doc, "What Worked Particularly Well", 2)
    add_list_item(doc, "Deep understanding of our financial domain logic â€” surpassed expectations for a QA engineer")
    add_list_item(doc, "Proactively proposed and implemented risk-based testing for payment authorisation flows")
    add_list_item(doc, "Excellent knowledge transfer: our internal QA team now maintains the framework independently")
    add_list_item(doc, "Fast ramp-up to internal toolchain: Jira, Confluence, Jenkins CI â€” productive from week two")

    add_heading(doc, "Areas for Improvement", 2)
    add_list_item(doc, "Visibility in steering committee meetings â€” sometimes defers when her input would add value")
    add_list_item(doc, "Executive-level summary reporting â€” more concise release-readiness summaries would be appreciated")

    add_heading(doc, "Additional Comments", 2)
    doc.add_paragraph(
        "Elena is one of the best QA engineers we have worked with. We have specifically requested "
        "her continued involvement in the 2026 programme extension.")

    doc.save("test_docs_comprehensive/13_Elena_Rodriguez_2025_ClientFeedback.docx")
    print("âœ“ Created: 13_Elena_Rodriguez_2025_ClientFeedback.docx")


def create_elena_project_feedback_2025():
    """Elena Rodriguez 2025 â€” Project Feedback (separate document)"""
    doc = Document()
    add_doc_header(doc, "PROJECT FEEDBACK", "FinCore Banking Platform â€” QA Automation Stream 2025")
    add_heading(doc, "Project Feedback Form", 1)
    add_heading(doc, "Review Period: January 2025 â€“ December 2025", 2)

    add_heading(doc, "Project Details", 2)
    doc.add_paragraph("Employee: Elena Rodriguez")
    doc.add_paragraph("Role on Project: Lead QA Automation Engineer")
    doc.add_paragraph("Project: FinCore Banking Platform â€” Automated Testing Stream")
    doc.add_paragraph("Project Duration: January 2025 â€“ December 2025")
    doc.add_paragraph("Reviewer: Samantha Okafor, Delivery Lead")

    add_section_type(doc, "PERFORMANCE METRICS")
    add_heading(doc, "Performance Ratings", 2)
    add_list_item(doc, "Technical contribution quality: 5 / 5")
    add_list_item(doc, "Collaboration and teamwork: 4 / 5")
    add_list_item(doc, "Meeting deadlines and sprint commitments: 5 / 5")
    add_list_item(doc, "Initiative and problem-solving: 5 / 5")
    add_list_item(doc, "Quality and maintainability of output: 5 / 5")
    add_list_item(doc, "Overall project contribution rating: 5 / 5")

    add_section_type(doc, "KEY ACHIEVEMENTS")
    add_heading(doc, "Notable Achievements on Project", 2)
    add_list_item(doc, "Designed and implemented 1,200+ automated test cases covering core banking workflows")
    add_list_item(doc, "Regression coverage: increased from 40% at project start to 95% by Q4")
    add_list_item(doc, "Integrated test suite into Jenkins CI/CD pipeline â€” release cycle reduced from 3 weeks to 1 week")
    add_list_item(doc, "Identified 3 critical security-related defects in payment authorisation module before go-live")
    add_list_item(doc, "Mentored 2 junior QA engineers: both now maintain framework features independently")

    add_section_type(doc, "FEEDBACK")
    add_heading(doc, "Technical Contribution", 2)
    add_feedback(doc, "Samantha Okafor, Delivery Lead",
        "Elena brought structured engineering discipline to our testing â€” proper design patterns, "
        "clear documentation, CI integration. The framework she built is genuinely production-grade. "
        "She is the most technically capable QA engineer I have managed.")

    add_heading(doc, "Collaboration and Teamwork", 2)
    add_feedback(doc, "Samantha Okafor, Delivery Lead",
        "Elena works well within the squad and the dev team respect her technical judgement. "
        "One development area: she sometimes works too independently and could involve the team earlier "
        "in design decisions so they feel greater ownership. Minor point overall â€” collaboration is strong.")
    add_feedback(doc, "Marcus Li, Senior Backend Developer",
        "Elena's test coverage found bugs I was convinced were correct. Her precision is unmatched. "
        "She explains defects clearly and without blame â€” genuinely constructive to work with.")

    add_heading(doc, "Initiative and Problem-Solving", 2)
    add_feedback(doc, "Samantha Okafor, Delivery Lead",
        "When the client changed requirements mid-sprint on the payment modules, Elena self-organised "
        "to refactor the affected test suites within two days without escalation. That kind of proactivity "
        "is exactly what we need on a critical financial project.")

    add_heading(doc, "Areas for Development", 2)
    add_list_item(doc, "Stakeholder communication: Strong technically but executive-level clarity still developing")
    add_list_item(doc, "Delegation: Tends to retain tasks personally; building multiplication skills would help her scale to lead")
    add_list_item(doc, "Cross-domain input: Has the capability but does not always volunteer opinions outside QA scope")

    doc.save("test_docs_comprehensive/13_Elena_Rodriguez_2025_ProjectFeedback.docx")
    print("âœ“ Created: 13_Elena_Rodriguez_2025_ProjectFeedback.docx")


def create_elena_pdp_2025():
    """Elena Rodriguez 2025 â€” Personal Development Plan (PDP)"""
    doc = Document()
    add_doc_header(doc, "PDP", "Personal Development Plan â€” 2025")
    add_heading(doc, "Personal Development Plan", 1)
    add_heading(doc, "Review Period: January 2025 â€“ December 2025", 2)

    add_heading(doc, "Employee Details", 2)
    doc.add_paragraph("Employee: Elena Rodriguez")
    doc.add_paragraph("Role: QA Automation Engineer (Mid-Senior)")
    doc.add_paragraph("Manager: Samantha Okafor")
    doc.add_paragraph("Department: Quality Engineering")
    doc.add_paragraph("Date Agreed: January 2025")

    add_section_type(doc, "SUMMARY")
    add_heading(doc, "PDP Summary", 2)
    doc.add_paragraph(
        "Elena is targeting the QA Tech Lead role in 2026. This PDP focuses on developing leadership "
        "and stakeholder communication skills alongside deepening technical expertise in performance "
        "and security testing. Technical automation skills are already advanced.")

    add_section_type(doc, "SKILLS & CERTIFICATION")
    add_heading(doc, "Current Skills Assessment", 2)
    add_list_item(doc, "Python test automation: Advanced (Pytest, Selenium, Playwright)")
    add_list_item(doc, "CI/CD integration: Proficient (Jenkins, GitHub Actions)")
    add_list_item(doc, "API testing: Advanced (Postman, RestAssured)")
    add_list_item(doc, "Test framework design: Strong")
    add_list_item(doc, "Defect analysis and root-cause investigation: Advanced")
    add_list_item(doc, "Financial services / banking domain knowledge: Solid")

    add_section_type(doc, "DEVELOPMENT AREAS")
    add_heading(doc, "Development Objectives â€” 2025", 2)

    add_heading(doc, "Objective 1: Stakeholder Communication and Influence", 3)
    doc.add_paragraph("Target: Become the confident QA voice in client steering meetings.")
    add_list_item(doc, "Lead QA status updates at steering committee level from Q2 2025")
    add_list_item(doc, "Prepare and present monthly quality dashboards to project sponsors")
    add_list_item(doc, "Complete Presentation and Executive Communication Workshop (Q1 2025)")
    add_list_item(doc, "Success metric: Client rates communication 4/5 or higher in Q4 review")

    add_heading(doc, "Objective 2: Performance and Security Testing", 3)
    doc.add_paragraph("Target: Expand testing capability into performance and security disciplines.")
    add_list_item(doc, "Complete ISTQB Advanced Level â€” Performance Testing (Q2 2025)")
    add_list_item(doc, "Complete OWASP Testing Guide self-study (Q3 2025)")
    add_list_item(doc, "Deliver and present performance test report on FinCore project (H1 2025)")
    add_list_item(doc, "Success metric: Performance test report accepted by client; ISTQB Advanced cert achieved")

    add_heading(doc, "Objective 3: Mentoring and Team Leadership", 3)
    doc.add_paragraph("Target: Develop others to prepare for transition to QA Lead.")
    add_list_item(doc, "Formally mentor 2 junior QA engineers throughout 2025 (monthly structured check-ins)")
    add_list_item(doc, "Lead team bi-weekly QA knowledge-share sessions")
    add_list_item(doc, "Co-design onboarding plan for new QA joiners (Q2)")
    add_list_item(doc, "Success metric: Both mentees complete individual certification goals by Q4")

    add_section_type(doc, "GOALS")
    add_heading(doc, "Career Goals", 2)

    add_heading(doc, "1-Year Goal (2025)", 3)
    doc.add_paragraph(
        "Establish technical authority on FinCore engagement. Build credibility as QA client lead. "
        "Achieve ISTQB Advanced â€” Performance Testing certification. Demonstrate mentoring capability.")

    add_heading(doc, "3-Year Goal (2028)", 3)
    doc.add_paragraph(
        "Reach QA Tech Lead / Principal Automation Engineer level. Own QA framework strategy across "
        "a programme or account. Lead a team of 4â€“6 engineers technically.")

    add_section_type(doc, "PDP - PROFESSIONAL DEVELOPMENT")
    add_heading(doc, "Planned Training and Certifications â€” 2025", 2)
    add_list_item(doc, "ISTQB Advanced Level â€” Performance Testing (Q2 2025)")
    add_list_item(doc, "Presentation and Executive Communication Workshop â€” external (Q1 2025)")
    add_list_item(doc, "OWASP Testing Guide â€” self-study (Q3 2025)")
    add_list_item(doc, "Python Advanced Patterns for Test Automation â€” online course (Q1 2025)")
    add_list_item(doc, "Internal Leadership Foundations Programme (Q3 2025)")

    add_heading(doc, "Manager Support Agreed", 2)
    add_list_item(doc, "Monthly PDP check-in in 1-on-1s")
    add_list_item(doc, "Protected learning time: 1 day per month")
    add_list_item(doc, "Stretch assignment: Lead QA technical governance on FinCore from Q2")
    add_list_item(doc, "Introductions to senior client stakeholders to build visibility")

    doc.save("test_docs_comprehensive/13_Elena_Rodriguez_2025_PDP.docx")
    print("âœ“ Created: 13_Elena_Rodriguez_2025_PDP.docx")


# ===== SCENARIO 14: JAMES PARK â€” SENIOR FULL-STACK DEVELOPER =====

def create_james_client_feedback_2025():
    """James Park 2025 â€” Client Feedback (separate document)"""
    doc = Document()
    add_doc_header(doc, "CLIENT FEEDBACK", "Senior Full-Stack Developer â€” NexaBank Digital Onboarding")
    add_heading(doc, "Client Feedback Form", 1)
    add_heading(doc, "Review Period: March 2025 â€“ December 2025", 2)

    add_heading(doc, "Engagement Details", 2)
    doc.add_paragraph("Employee: James Park")
    doc.add_paragraph("Role: Senior Full-Stack Developer / Technical Design Lead")
    doc.add_paragraph("Project / Engagement: NexaBank â€” Digital Onboarding Platform (Phase 1)")
    doc.add_paragraph("Client Organisation: NexaBank plc")
    doc.add_paragraph("Client Contact: Diana Chen, CTO, NexaBank plc")

    add_section_type(doc, "CLIENT FEEDBACK")
    add_heading(doc, "Satisfaction Ratings", 2)
    add_list_item(doc, "Overall satisfaction: 5 / 5")
    add_list_item(doc, "Technical competence and architecture quality: 5 / 5")
    add_list_item(doc, "Communication and workshop facilitation: 5 / 5")
    add_list_item(doc, "Collaboration with client team: 5 / 5")
    add_list_item(doc, "Meeting delivery commitments: 5 / 5")
    add_list_item(doc, "Responsiveness to change requests: 4 / 5")
    add_list_item(doc, "Would you engage this person again: Yes â€” strongly and without hesitation")

    add_heading(doc, "Technical Quality and Architecture", 2)
    add_feedback(doc, "Diana Chen, CTO â€” NexaBank plc",
        "James is outstanding. The onboarding platform architecture he proposed was more elegant and "
        "scalable than what our internal team had planned. He brought genuine expertise to every technical "
        "decision and was never afraid to challenge our assumptions constructively. His API design work "
        "has set a new standard we are rolling out across other systems.")

    add_heading(doc, "Communication and Workshop Facilitation", 2)
    add_feedback(doc, "Diana Chen, CTO â€” NexaBank plc",
        "James ran three architecture workshops with our leadership team. He has a rare ability to make "
        "complex technical topics accessible to non-technical executives. Several of our board-level "
        "stakeholders commented specifically on how clearly he explained trade-offs.")
    add_feedback(doc, "Tom Vickers, Head of Product â€” NexaBank plc",
        "James bridged the gap between engineering and product perfectly. When we had conflicting "
        "priorities he helped us make principled decisions rather than reactive ones. I would describe "
        "him as a trusted advisor, not just a developer.")

    add_heading(doc, "What Worked Particularly Well", 2)
    add_list_item(doc, "Architecture documentation: Comprehensive and accessible â€” raised our internal engineering standards")
    add_list_item(doc, "API design: Clean, well-versioned, OpenAPI spec'd â€” frontend team had zero integration issues")
    add_list_item(doc, "Security consciousness: Proactively flagged PII handling risk before it became a compliance issue")
    add_list_item(doc, "Knowledge transfer: Genuinely upskilled three of our internal developers")

    add_heading(doc, "Areas for Improvement", 2)
    add_list_item(doc, "Written documentation occasionally overdetailed â€” executive summaries could be shorter")
    add_list_item(doc, "Could delegate more to junior team members; sometimes retains ownership of tasks that would develop the team")

    add_heading(doc, "Additional Comments", 2)
    doc.add_paragraph(
        "James has been one of the most impactful individuals we have engaged in five years. "
        "We have extended his contract through Q2 2026 and are in discussion about a permanent advisory arrangement.")

    doc.save("test_docs_comprehensive/14_James_Park_2025_ClientFeedback.docx")
    print("âœ“ Created: 14_James_Park_2025_ClientFeedback.docx")


def create_james_project_feedback_2025():
    """James Park 2025 â€” Project Feedback (separate document)"""
    doc = Document()
    add_doc_header(doc, "PROJECT FEEDBACK", "NexaBank Digital Onboarding Platform â€” 2025")
    add_heading(doc, "Project Feedback Form", 1)
    add_heading(doc, "Review Period: March 2025 â€“ December 2025", 2)

    add_heading(doc, "Project Details", 2)
    doc.add_paragraph("Employee: James Park")
    doc.add_paragraph("Role on Project: Senior Full-Stack Developer / Technical Design Lead")
    doc.add_paragraph("Project: NexaBank â€” Digital Onboarding Platform (Phase 1)")
    doc.add_paragraph("Project Duration: March 2025 â€“ December 2025")
    doc.add_paragraph("Reviewer: Natalie Brooks, Programme Manager")

    add_section_type(doc, "PERFORMANCE METRICS")
    add_heading(doc, "Performance Ratings", 2)
    add_list_item(doc, "Technical contribution quality: 5 / 5")
    add_list_item(doc, "Architecture and technical design leadership: 5 / 5")
    add_list_item(doc, "Collaboration and teamwork: 5 / 5")
    add_list_item(doc, "Code quality and standards: 5 / 5")
    add_list_item(doc, "Delivery against sprint commitments: 4 / 5")
    add_list_item(doc, "Overall project contribution rating: 5 / 5")

    add_section_type(doc, "KEY ACHIEVEMENTS")
    add_heading(doc, "Notable Achievements on Project", 2)
    add_list_item(doc, "Designed microservices architecture: 5 services, 3 integration layers, full OpenAPI specification")
    add_list_item(doc, "Delivered end-to-end onboarding flow: KYC verification, account creation, document upload")
    add_list_item(doc, "End-to-end response time: < 2 seconds (requirement was < 3 seconds)")
    add_list_item(doc, "Zero critical security findings in third-party penetration test â€” first time in 3 years for this client")
    add_list_item(doc, "Introduced contract-first API design approach â€” adopted as team and then practice standard")
    add_list_item(doc, "Mentored 3 client-side developers; all rated the mentoring highly valuable in feedback")

    add_section_type(doc, "FEEDBACK")
    add_heading(doc, "Technical Leadership and Architecture", 2)
    add_feedback(doc, "Natalie Brooks, Programme Manager",
        "James took ownership of technical design in a way I had not anticipated from a developer "
        "rather than an architect. He proactively scheduled architecture review sessions, prepared "
        "crisp decision documents, and coached the team on implementation patterns. It freed me to "
        "focus on delivery management â€” that is exactly the partnership I need on a complex programme.")
    add_feedback(doc, "Sofia Martinez, Backend Developer",
        "James's code reviews are the best I have experienced. He explains the reasoning, not just "
        "the fix. My skills have measurably improved over nine months of working alongside him.")

    add_heading(doc, "Delivery", 2)
    add_feedback(doc, "Natalie Brooks, Programme Manager",
        "Phase 1 delivered on schedule. One sprint was at risk due to a third-party API issue; James "
        "found a workaround within 48 hours. His estimation accuracy is good overall â€” occasionally "
        "optimistic on tasks with external dependencies, but he always flags the risk early.")

    add_heading(doc, "Areas for Development", 2)
    add_list_item(doc, "Estimation: Slight optimism bias on externally-dependent tasks â€” recommend adding buffer in planning")
    add_list_item(doc, "Delegation: Strongest performer on team but should involve mid-level engineers more in design choices")
    add_list_item(doc, "Documentation maintenance: Architecture docs excellent at sprint start; not always updated after mid-sprint changes")

    doc.save("test_docs_comprehensive/14_James_Park_2025_ProjectFeedback.docx")
    print("âœ“ Created: 14_James_Park_2025_ProjectFeedback.docx")


def create_james_pdp_2025():
    """James Park 2025 â€” Personal Development Plan (PDP)"""
    doc = Document()
    add_doc_header(doc, "PDP", "Personal Development Plan â€” 2025")
    add_heading(doc, "Personal Development Plan", 1)
    add_heading(doc, "Review Period: January 2025 â€“ December 2025", 2)

    add_heading(doc, "Employee Details", 2)
    doc.add_paragraph("Employee: James Park")
    doc.add_paragraph("Role: Senior Full-Stack Developer")
    doc.add_paragraph("Manager: Natalie Brooks")
    doc.add_paragraph("Department: Software Engineering â€” Financial Services Practice")
    doc.add_paragraph("Date Agreed: January 2025")

    add_section_type(doc, "SUMMARY")
    add_heading(doc, "PDP Summary", 2)
    doc.add_paragraph(
        "James is targeting the Staff Engineer level by end of 2026. This PDP focuses on formalising "
        "architecture leadership through accreditation, developing structured team-development skills, "
        "and building external industry visibility. Core technical skills are already at expert level.")

    add_section_type(doc, "SKILLS & CERTIFICATION")
    add_heading(doc, "Current Skills Assessment", 2)
    add_list_item(doc, "Full-stack development: Expert (React / TypeScript, Node.js, Python, Go)")
    add_list_item(doc, "API design: Expert (REST, GraphQL, OpenAPI contract-first)")
    add_list_item(doc, "Microservices architecture: Advanced")
    add_list_item(doc, "Cloud (AWS): Advanced â€” EC2, RDS, Lambda, SQS, S3, API Gateway")
    add_list_item(doc, "Security-aware development: Advanced (OWASP, PCI DSS awareness)")
    add_list_item(doc, "Client-facing communication and facilitation: Strong")
    add_list_item(doc, "Mentoring: Strong (informal; being formalised this year)")

    add_section_type(doc, "DEVELOPMENT AREAS")
    add_heading(doc, "Development Objectives â€” 2025", 2)

    add_heading(doc, "Objective 1: Formal Architecture Accreditation", 3)
    doc.add_paragraph("Target: Achieve a recognised solution architecture certification to formalise expertise.")
    add_list_item(doc, "AWS Solutions Architect â€” Professional (target Q2 2025)")
    add_list_item(doc, "TOGAF 10 Foundation self-study (Q3 2025, optional exam Q4)")
    add_list_item(doc, "Lead at least 2 formal Architecture Decision Record (ADR) processes on live engagements")
    add_list_item(doc, "Success metric: AWS SA Professional certificate achieved by June 2025; 2 ADRs approved by client")

    add_heading(doc, "Objective 2: Engineering Leadership and Structured Delegation", 3)
    doc.add_paragraph("Target: Build structured team-development practice to prepare for Staff Engineer scope.")
    add_list_item(doc, "Formally mentor 2 junior developers â€” monthly structured check-ins with written goals")
    add_list_item(doc, "Run 4 internal architecture review sessions across the year")
    add_list_item(doc, "Consciously delegate implementation tasks to mid-level engineers with structured coaching")
    add_list_item(doc, "Success metric: Both mentees complete individual development goals; delivery manager notes delegation improvement in Q4")

    add_heading(doc, "Objective 3: Industry Visibility", 3)
    doc.add_paragraph("Target: Build external profile to support Staff Engineer and future principal track progression.")
    add_list_item(doc, "Publish 2 technical articles on microservices in regulated industries")
    add_list_item(doc, "Submit a conference talk proposal to at least 1 engineering conference")
    add_list_item(doc, "Make a meaningful contribution to 1 open-source project")
    add_list_item(doc, "Success metric: 2 articles published; conference submission made by Q4")

    add_section_type(doc, "GOALS")
    add_heading(doc, "Career Goals", 2)

    add_heading(doc, "1-Year Goal (2025)", 3)
    doc.add_paragraph(
        "Achieve AWS Solutions Architect â€” Professional certification. Lead technical design on "
        "NexaBank Phase 2 from April. Be recognised by client as primary technical authority. "
        "Begin formal mentoring programme with 2 developers.")

    add_heading(doc, "3-Year Goal (2028)", 3)
    doc.add_paragraph(
        "Reach Staff Engineer / Principal Engineer level within the practice. Lead technical strategy "
        "across multi-team engagements. Develop a recognised specialism in fintech and regulated-industry "
        "system architecture.")

    add_section_type(doc, "PDP - PROFESSIONAL DEVELOPMENT")
    add_heading(doc, "Planned Training and Certifications â€” 2025", 2)
    add_list_item(doc, "AWS Solutions Architect â€” Professional (Q2 2025)")
    add_list_item(doc, "TOGAF 10 Foundation (Q3 2025)")
    add_list_item(doc, "Engineering Leadership Fundamentals â€” internal programme (Q1 2025)")
    add_list_item(doc, "Advanced System Design Workshop â€” external (Q2 2025)")
    add_list_item(doc, "PCI DSS Developer Awareness â€” required for NexaBank project (Q1 2025)")

    add_heading(doc, "Manager Support Agreed", 2)
    add_list_item(doc, "Dedicated 10% time allocation for learning and certification preparation")
    add_list_item(doc, "Formal mentoring pairing with 2 junior developers from Q1")
    add_list_item(doc, "Sponsorship and support for conference talk submission")
    add_list_item(doc, "Quarterly PDP review in 1-on-1s")

    doc.save("test_docs_comprehensive/14_James_Park_2025_PDP.docx")
    print("âœ“ Created: 14_James_Park_2025_PDP.docx")


# ===== SCENARIO 15: PRIYA SHARMA â€” QA LEAD, FIRST YEAR =====

def create_priya_client_feedback_2025():
    """Priya Sharma 2025 â€” Client Feedback (separate document)"""
    doc = Document()
    add_doc_header(doc, "CLIENT FEEDBACK", "QA Lead â€” Covanta InsureTech Policy Management")
    add_heading(doc, "Client Feedback Form", 1)
    add_heading(doc, "Review Period: April 2025 â€“ December 2025", 2)

    add_heading(doc, "Engagement Details", 2)
    doc.add_paragraph("Employee: Priya Sharma")
    doc.add_paragraph("Role: QA Lead")
    doc.add_paragraph("Project / Engagement: Covanta InsureTech â€” Policy Management System Modernisation")
    doc.add_paragraph("Client Organisation: Covanta InsureTech Ltd")
    doc.add_paragraph("Client Contact: Robert Inglesby, Head of Delivery, Covanta InsureTech")
    doc.add_paragraph("Note: Priya joined this project in April 2025 (appointed to QA Lead role January 2025)")

    add_section_type(doc, "CLIENT FEEDBACK")
    add_heading(doc, "Satisfaction Ratings", 2)
    add_list_item(doc, "Overall satisfaction: 4 / 5")
    add_list_item(doc, "Technical QA competence: 4 / 5")
    add_list_item(doc, "Communication and reporting: 3 / 5")
    add_list_item(doc, "Team management and leadership visibility: 3 / 5")
    add_list_item(doc, "Responsiveness: 4 / 5")
    add_list_item(doc, "Meeting delivery expectations: 4 / 5")
    add_list_item(doc, "Would you engage this person again: Yes")

    add_heading(doc, "Technical Quality and Testing Approach", 2)
    add_feedback(doc, "Robert Inglesby, Head of Delivery â€” Covanta InsureTech",
        "Priya has strong technical QA skills. Her test planning is thorough and she clearly understands "
        "our insurance domain. The regression suite she built for the policy engine performs well and her "
        "defect triage is fast and accurate. She demonstrates sound engineering judgement.")

    add_heading(doc, "Communication and Reporting", 2)
    add_feedback(doc, "Robert Inglesby, Head of Delivery â€” Covanta InsureTech",
        "This is an area for development. Priya is strong in technical discussions but the quality "
        "reporting provided to our stakeholders lacks executive clarity. We sometimes need follow-up "
        "questions to understand release readiness. We would welcome more concise, risk-focused summaries.")
    add_feedback(doc, "Sandra Patel, QA Manager â€” Covanta InsureTech",
        "Priya is very knowledgeable â€” one-to-one communication is great. In group settings, especially "
        "with senior stakeholders, she could be more assertive. On several occasions her team's good work "
        "was not sufficiently visible at programme level.")

    add_heading(doc, "Team Leadership", 2)
    add_feedback(doc, "Robert Inglesby, Head of Delivery â€” Covanta InsureTech",
        "Priya manages her QA team well at a day-to-day task level and is well regarded by her engineers. "
        "We would like to see more strategic thinking on process improvement and testing roadmap â€” "
        "that fuller picture view is what we expect from someone in the QA Lead role.")

    add_heading(doc, "What Worked Well", 2)
    add_list_item(doc, "Deep insurance domain knowledge â€” accelerated test design significantly from day one")
    add_list_item(doc, "Team morale: Priya's QA team is clearly motivated and performing well")
    add_list_item(doc, "Defect triage: Fast and accurate â€” no false positives under her lead")
    add_list_item(doc, "Dev collaboration: Strong, non-confrontational working relationship with engineering team")

    add_heading(doc, "Areas for Improvement", 2)
    add_list_item(doc, "Executive-level reporting: Needs to be more concise, risk-focused, and audience-aware")
    add_list_item(doc, "Strategic QA roadmap ownership: Should propose process improvements proactively")
    add_list_item(doc, "Confidence in senior meetings: Has the knowledge â€” needs to use her voice more")

    doc.save("test_docs_comprehensive/15_Priya_Sharma_2025_ClientFeedback.docx")
    print("âœ“ Created: 15_Priya_Sharma_2025_ClientFeedback.docx")


def create_priya_project_feedback_2025():
    """Priya Sharma 2025 â€” Project Feedback (separate document)"""
    doc = Document()
    add_doc_header(doc, "PROJECT FEEDBACK", "Covanta InsureTech â€” Policy Management System 2025")
    add_heading(doc, "Project Feedback Form", 1)
    add_heading(doc, "Review Period: April 2025 â€“ December 2025", 2)

    add_heading(doc, "Project Details", 2)
    doc.add_paragraph("Employee: Priya Sharma")
    doc.add_paragraph("Role on Project: QA Lead")
    doc.add_paragraph("Project: Covanta InsureTech â€” Policy Management System Modernisation (Phase 1)")
    doc.add_paragraph("Project Duration: April 2025 â€“ December 2025")
    doc.add_paragraph("Reviewer: Ahmed Khalil, Programme Director")
    doc.add_paragraph("Note: First engagement as QA Lead â€” previously QA Engineer on prior projects")

    add_section_type(doc, "PERFORMANCE METRICS")
    add_heading(doc, "Performance Ratings", 2)
    add_list_item(doc, "Technical QA quality: 4 / 5")
    add_list_item(doc, "Test strategy ownership: 4 / 5")
    add_list_item(doc, "Team leadership and management: 3 / 5")
    add_list_item(doc, "Stakeholder communication: 3 / 5")
    add_list_item(doc, "Delivery against QA milestones: 4 / 5")
    add_list_item(doc, "Overall project contribution rating: 4 / 5")

    add_section_type(doc, "KEY ACHIEVEMENTS")
    add_heading(doc, "Notable Achievements on Project", 2)
    add_list_item(doc, "Designed test strategy for a 6-module policy management system from scratch")
    add_list_item(doc, "Built and led a QA team of 4 engineers across onshore / offshore model")
    add_list_item(doc, "Delivered Phase 1 regression suite: 800 test cases, 88% automated")
    add_list_item(doc, "Zero escaped defects to UAT in final 2 months â€” improved from 6 escapes in first quarter")
    add_list_item(doc, "Introduced defect-prevention metrics to team â€” adopted as internal project standard")
    add_list_item(doc, "Managed critical delivery risk when a QA resource left mid-project without disrupting timeline")

    add_section_type(doc, "FEEDBACK")
    add_heading(doc, "Test Strategy and Technical Leadership", 2)
    add_feedback(doc, "Ahmed Khalil, Programme Director",
        "Priya designed a mature, risk-based test strategy for a complex insurance system. "
        "Her domain knowledge is impressive and she made defensible, well-reasoned decisions about "
        "coverage priorities that stood up to client challenge. That is not easy for someone in their "
        "first QA Lead engagement.")
    add_feedback(doc, "Omar Reyes, Technical Architect",
        "Priya knows her craft. Defect reports are detailed and accurate. She caught several integration "
        "issues that would have been costly post-launch. Strong technical instincts.")

    add_heading(doc, "Team Leadership", 2)
    add_feedback(doc, "Ahmed Khalil, Programme Director",
        "This is Priya's first time leading a QA team and the growth has been visible across the nine months. "
        "She started tentatively â€” checking in frequently for reassurance â€” and finished Q4 making confident "
        "decisions independently. Her team speak highly of her. Delegation and workload distribution are "
        "the main areas still to develop.")
    add_feedback(doc, "Yuki Tanaka, QA Engineer (reporting to Priya)",
        "Priya is a great lead to work for. She explains the why, not just the what. She is still new to "
        "management but she is learning fast and is genuinely supportive of her team.")

    add_heading(doc, "Areas for Development", 2)
    add_list_item(doc, "Structured delegation: Learning to allocate ownership clearly without retaining all technical tasks herself")
    add_list_item(doc, "Risk escalation timing: Tends to attempt internal resolution before escalating â€” earlier escalation would help programme visibility")
    add_list_item(doc, "Programme-level reporting: Detailed QA reports well-received by engineering; senior stakeholder summaries still improving")
    add_list_item(doc, "Confidence in programme governance forums: Contributions are good but visibility could be higher")

    doc.save("test_docs_comprehensive/15_Priya_Sharma_2025_ProjectFeedback.docx")
    print("âœ“ Created: 15_Priya_Sharma_2025_ProjectFeedback.docx")


def create_priya_pdp_2025():
    """Priya Sharma 2025 â€” Personal Development Plan (PDP)"""
    doc = Document()
    add_doc_header(doc, "PDP", "Personal Development Plan â€” 2025")
    add_heading(doc, "Personal Development Plan", 1)
    add_heading(doc, "Review Period: January 2025 â€“ December 2025", 2)

    add_heading(doc, "Employee Details", 2)
    doc.add_paragraph("Employee: Priya Sharma")
    doc.add_paragraph("Role: QA Lead (promoted January 2025 â€” first year in role)")
    doc.add_paragraph("Manager: Ahmed Khalil")
    doc.add_paragraph("Department: Quality Engineering")
    doc.add_paragraph("Date Agreed: January 2025")

    add_section_type(doc, "SUMMARY")
    add_heading(doc, "PDP Summary", 2)
    doc.add_paragraph(
        "Priya transitioned into the QA Lead role in January 2025 after 3 years as a QA Engineer. "
        "This PDP addresses the leadership, stakeholder communication, and strategic skills required "
        "to succeed in the lead role. Technical QA skills are already strong and are not the primary focus.")

    add_section_type(doc, "SKILLS & CERTIFICATION")
    add_heading(doc, "Current Skills Assessment", 2)
    add_list_item(doc, "Manual and exploratory testing: Expert")
    add_list_item(doc, "Test automation: Proficient (Selenium, Python, Pytest)")
    add_list_item(doc, "Test planning and strategy: Strong (recently developed at lead level)")
    add_list_item(doc, "Insurance and policy management domain knowledge: Strong")
    add_list_item(doc, "Defect analysis and prevention: Advanced")
    add_list_item(doc, "ISTQB Foundation Level: Certified (2023)")

    add_section_type(doc, "DEVELOPMENT AREAS")
    add_heading(doc, "Development Objectives â€” 2025", 2)

    add_heading(doc, "Objective 1: Leadership and Team Management", 3)
    doc.add_paragraph("Target: Build core people management skills for leading a team of four.")
    add_list_item(doc, "Complete Internal Leadership Foundations Programme (Q2 2025)")
    add_list_item(doc, "Hold structured 1-on-1s with all direct reports monthly from Q1")
    add_list_item(doc, "Complete Giving Developmental Feedback training (Q2 2025)")
    add_list_item(doc, "Practice structured delegation on Covanta project from Q2")
    add_list_item(doc, "Success metric: Team satisfaction score > 3.8 / 5 in Q4 team survey")

    add_heading(doc, "Objective 2: Stakeholder Communication and Reporting", 3)
    doc.add_paragraph("Target: Build confidence and skill communicating with senior client stakeholders.")
    add_list_item(doc, "Complete Communication for Technical Leaders Workshop (Q1 2025)")
    add_list_item(doc, "Own and present quality dashboard at all client steering meetings from Q2")
    add_list_item(doc, "Develop a QA reporting template designed for senior / executive audience")
    add_list_item(doc, "Success metric: Client rates communication 4 / 5 or higher in end-of-year feedback")

    add_heading(doc, "Objective 3: QA Strategy and Process Leadership", 3)
    doc.add_paragraph("Target: Move from executing QA strategy to owning and evolving it.")
    add_list_item(doc, "Publish internal QA Standards Guide for the practice (Q3 2025)")
    add_list_item(doc, "Complete ISTQB Advanced Level â€” Test Management (Q3 2025)")
    add_list_item(doc, "Propose and lead one quality improvement initiative on the current project")
    add_list_item(doc, "Success metric: ISTQB Advanced achieved; QA Standards Guide adopted on at least one other project")

    add_section_type(doc, "GOALS")
    add_heading(doc, "Career Goals", 2)

    add_heading(doc, "1-Year Goal (2025)", 3)
    doc.add_paragraph(
        "Successfully establish effectiveness and credibility in QA Lead role. Deliver Phase 1 "
        "on Covanta to high quality. Build confidence in senior stakeholder communication. "
        "Complete ISTQB Advanced â€” Test Management certification.")

    add_heading(doc, "3-Year Goal (2028)", 3)
    doc.add_paragraph(
        "Reach Senior QA Lead / QA Practice Lead level. Own QA strategy and standards across "
        "multiple projects or a programme. Lead a team of 6â€“8. Build toward QA Chapter Lead "
        "or Head of Quality Engineering path.")

    add_section_type(doc, "PDP - PROFESSIONAL DEVELOPMENT")
    add_heading(doc, "Planned Training and Certifications â€” 2025", 2)
    add_list_item(doc, "ISTQB Advanced Level â€” Test Management (Q3 2025)")
    add_list_item(doc, "Communication for Technical Leaders Workshop â€” external (Q1 2025)")
    add_list_item(doc, "Internal Leadership Foundations Programme (Q2 2025)")
    add_list_item(doc, "Giving Developmental Feedback Skills Training (Q2 2025)")
    add_list_item(doc, "Risk-Based Testing Masterclass â€” online self-study (Q4 2025)")

    add_heading(doc, "Manager Support Agreed", 2)
    add_list_item(doc, "Bi-weekly 1-on-1 PDP check-ins for first 6 months, then monthly")
    add_list_item(doc, "Shadow Ahmed in programme-level stakeholder meetings throughout Q1")
    add_list_item(doc, "Transition steering committee quality reporting ownership to Priya by Q2")
    add_list_item(doc, "Introduction to QA Lead peer group in the practice for knowledge sharing")

    doc.save("test_docs_comprehensive/15_Priya_Sharma_2025_PDP.docx")
    print("âœ“ Created: 15_Priya_Sharma_2025_PDP.docx")


def main():
    """Create all test documents"""
    print("\n" + "=" * 90)
    print(" TEST DATASET GENERATION")
    print("=" * 90 + "\n")

    print("MULTI-DOCUMENT PERSONAS (CLIENT FEEDBACK + PROJECT FEEDBACK + PDP)")
    print("â”€" * 90)
    print("PERSONA 13: ELENA RODRIGUEZ â€” QA Automation Engineer, FinCore Banking Platform")
    print("â”€" * 90)
    create_elena_client_feedback_2025()
    create_elena_project_feedback_2025()
    create_elena_pdp_2025()

    print("\nPERSONA 14: JAMES PARK â€” Senior Full-Stack Developer, NexaBank Digital Onboarding")
    print("â”€" * 90)
    create_james_client_feedback_2025()
    create_james_project_feedback_2025()
    create_james_pdp_2025()

    print("\nPERSONA 15: PRIYA SHARMA â€” QA Lead (First Year), Covanta InsureTech")
    print("â”€" * 90)
    create_priya_client_feedback_2025()
    create_priya_project_feedback_2025()
    create_priya_pdp_2025()

    print("\n" + "=" * 90)
    print(" TEST DATA GENERATION COMPLETE")
    print("=" * 90)
    
    print("\n Generated Documents: test_docs_comprehensive/")
    print("\n Personas:")
    print("   13. Elena Rodriguez  â”‚ 2025  â”‚ QA Automation Engineer (Client FB + Project FB + PDP)")
    print("   14. James Park       â”‚ 2025  â”‚ Senior Full-Stack Developer (Client FB + Project FB + PDP)")
    print("   15. Priya Sharma     â”‚ 2025  â”‚ QA Lead Year 1 (Client FB + Project FB + PDP)")
    print("\n Total Documents: 9 (3 personas x 3 document types)\n")


if __name__ == "__main__":
    main()
