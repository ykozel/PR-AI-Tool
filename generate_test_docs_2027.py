"""
Script to generate test .docx files for 2027 (Emma Laurent).

Persona: Emma Laurent, Senior Business Analyst, Strategy & Innovation dept.
These documents test the year linking feature - they reference 2025/2026 achievements
and build on the previous development plan.

Run:
    pip install python-docx
    python generate_test_docs_2027.py

This generates:
    1. Client Feedback for 2027
    2. PDP for 2027/2028
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ----- Helper Functions -----

def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)
    return p


def _body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def _meta_table(doc: Document, name: str, dept: str, role: str, period: str):
    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    labels = [
        ("Employee Name:", name, "Period:", period),
        ("Department:", dept, "Role:", role),
        ("Review Date:", "March 2027", "", ""),
    ]
    for row_idx, (l1, v1, l2, v2) in enumerate(labels):
        row = table.rows[row_idx]
        row.cells[0].text = l1
        row.cells[1].text = v1
        row.cells[2].text = l2
        row.cells[3].text = v2
    doc.add_paragraph()  # spacer


# ========================================
# CLIENT FEEDBACK - 2027
# ========================================

def create_client_feedback_2027(path: str):
    """
    Client feedback for Emma Laurent - 2027
    (Feedback from client stakeholder on performance in recent engagement)
    """
    doc = Document()
    doc.add_heading("CLIENT FEEDBACK FORM - 2027 ANNUAL REVIEW", 0)
    _meta_table(
        doc,
        name="Emma Laurent",
        dept="Strategy & Innovation",
        role="Senior Business Analyst",
        period="January - December 2027",
    )

    _heading(doc, "Client / Stakeholder Information", 1)
    client_info = doc.add_table(rows=4, cols=2)
    client_info.style = "Table Grid"
    client_rows = [
        ("Client Organization:", "Meridian Retail International Limited"),
        ("Project Name:", "Omnichannel Inventory Platform - Phase 2"),
        ("Stakeholder Name & Title:", "Rebecca Martinez, Head of UK Supply Chain Operations"),
        ("Feedback Date:", "January 15, 2027"),
    ]
    for idx, (label, value) in enumerate(client_rows):
        row = client_info.rows[idx]
        row.cells[0].text = label
        row.cells[1].text = value
    doc.add_paragraph()
    doc.add_paragraph()

    _heading(doc, "1. Business Analysis and Elicitation", 1)
    _body(doc,
        "Rating: Exceeds Expectations\n\n"
        "Emma's requirements elicitation was exceptionally strong. She conducted structured "
        "workshops with 15+ stakeholders across supply chain, retail operations, and IT teams, "
        "and synthesized competing requirements into a unified backlog that the delivery team "
        "executed without significant rework.\n\n"
        "Key strengths:\n"
        "  • Identified and resolved three critical scope conflicts before they impacted delivery\n"
        "  • Authored rigorous UAT test cases (87 test scripts) that caught 23 build defects pre-release\n"
        "  • Proactively documented process-improvement recommendations beyond the project scope\n"
        "  • Requirements traceability matrix was instrumental in change-request triage"
    )

    _heading(doc, "2. Communication and Stakeholder Management", 1)
    _body(doc,
        "Rating: Exceeds Expectations\n\n"
        "Emma's communication approach gave stakeholders confidence throughout the 14-month engagement. "
        "She owned the weekly steering group update, presented clearly under time pressure, and managed "
        "competing priorities without deflecting accountability.\n\n"
        "Standout moments:\n"
        "  • Presented Phase 2 investment case to the Meridian IT board (12 executives); the proposal passed with unanimous approval\n"
        "  • Facilitated a contentious scope negotiation between retail and supply chain teams; both groups signed off within one meeting\n"
        "  • Proactively escalated resource conflicts 4 weeks in advance, enabling mitigation planning\n"
        "  • Post-go-live debrief report was distributed to 30+ executives; formed the basis for Phase 3 planning"
    )

    _heading(doc, "3. Domain Knowledge and Business Acumen", 1)
    _body(doc,
        "Rating: Meets Expectations\n\n"
        "Emma demonstrated solid understanding of retail supply chain operations and quickly grasped "
        "Meridian's complex multi-warehouse inventory model. By mid-project, she was asking insightful "
        "questions about edge cases and business variability.\n\n"
        "Areas of strength:\n"
        "  • Learned Meridian's Warehouse Management System (WMS) architecture from first principles\n"
        "  • Identified system integration requirements that the internal IT team had initially overlooked\n"
        "  • Validated UAT scenarios against actual supply chain processes (good mastery by project end)\n\n"
        "Note: Her retail domain knowledge is developing. For a future large-scale retail engagement, "
        "6 weeks of pre-project immersion would further accelerate her contributions."
    )

    _heading(doc, "4. Problem-Solving and Decision Support", 1)
    _body(doc,
        "Rating: Exceeds Expectations\n\n"
        "Emma was instrumental in resolving three significant project challenges:\n\n"
        "  Challenge 1 - Reporting accuracy: Phase 1 reporting did not reconcile with GL systems.\n"
        "    Emma's action: Reverse-engineered the transactional data flow, identified a missing data "
        "  enrichment step, and worked with the data engineer to implement a fix.\n\n"
        "  Challenge 2 - Stakeholder deadlock on inventory-aging policy.\n"
        "    Emma's action: Facilitated a workshop comparing three policy options with financial models "
        "  showing the trade-off between write-off risk and cash flow. The group adopted a hybrid policy "
        "  that emerged from her structured analysis.\n\n"
        "  Challenge 3 - UAT schedule compression.\n"
        "    Emma's action: Risk-prioritized the 200+ planned test scenarios into 87 core scenarios, "
        "  removing redundant coverage. UAT completed on schedule without sacrificing quality."
    )

    _heading(doc, "5. Project Delivery Impact", 1)
    _body(doc,
        "Rating: Exceeds Expectations\n\n"
        "Emma's contributions directly enabled on-time, on-budget delivery:\n\n"
        "  • Requirements rework was <5%, compared to industry average of 12-15% for comparable engagements\n"
        "  • UAT defect detection rate: 23 build defects (vs. 8 typical for Phase 2 scope)\n"
        "  • Project delivered 2 weeks early; Emma's scope clarity contributed significantly to this acceleration\n"
        "  • Post-go-live support required <40 hours in first 30 days (vs. average of 80-120 hours for similar systems)\n"
        "  • Meridian approved Phase 3 funding immediately following Phase 2 close; Phase 3 scope includes "
        "    recommendations from Emma's post-go-live report"
    )

    _heading(doc, "6. Teamwork and Collaboration", 1)
    _body(doc,
        "Rating: Exceeds Expectations\n\n"
        "Emma was a collaborative and supportive team member throughout the engagement.\n\n"
        "  • Seamlessly integrated with the 8-person delivery team; no friction or role confusion\n"
        "  • Proactively documented requirements patterns and UAT best practices for team knowledge-sharing\n"
        "  • Supported the project manager in resource planning and risk identification\n"
        "  • Provided mentorship to a junior BA (David Chen) on requirements techniques; David's solo requirements "
        "    sessions in Phase 2 were of high quality"
    )

    _heading(doc, "7. Innovation & Continuous Improvement", 1)
    _body(doc,
        "Rating: Exceeds Expectations\n\n"
        "Beyond core BA responsibilities, Emma suggested several innovations:\n\n"
        "  • Proposed a client-side requirements dashboard (Power BI) for real-time backlog visibility - adopted by Meridian's internal team\n"
        "  • Identified and documented three opportunities for process automation, one implemented in Phase 2\n"
        "  • Recommended a post-implementation business review framework; Meridian is adopting this for Phase 3\n"
        "  • Shared learnings with the wider practice (invited to present at internal BA Community of Practice)"
    )

    _heading(doc, "8. Overall Assessment & Recommendation", 1)
    _body(doc,
        "Overall Rating: Exceeds Expectations\n\n"
        "Emma is a trusted, highly valued member of the Meridian engagement delivery team. She demonstrated "
        "exceptional requirements elicitation, stakeholder management, and problem-solving throughout Phase 2. "
        "Her contributions were instrumental in achieving on-time, on-budget delivery and strong stakeholder "
        "satisfaction.\n\n"
        "Meridian has requested Emma specifically for Phase 3 (2027-2028). We highly recommend continuing "
        "her engagement on the account.\n\n"
        "For promotion considerations: Emma has demonstrated senior-level competency in this engagement and "
        "should be considered for lead BA responsibilities on future account work.\n\n"
        "Signed: Rebecca Martinez, Head of UK Supply Chain Operations\n"
        "Date: January 15, 2027"
    )

    doc.save(path)
    print(f"[OK] {path}")


# ========================================
# PERSONAL DEVELOPMENT PLAN - 2027/2028
# ========================================

def create_pdp_2027_2028(path: str):
    """
    Personal Development Plan for Emma Laurent - 2027/2028 cycle.
    Builds on 2025/2026 PDP; includes progress check-in and new goals.
    """
    doc = Document()
    doc.add_heading("PERSONAL DEVELOPMENT PLAN - 2027 / 2028", 0)

    # Meta table
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    meta = [
        ("Employee Name:", "Emma Laurent", "Employee ID:", "EMP-2024-007832"),
        ("Department:",    "Strategy & Innovation", "Role:",   "Senior Business Analyst"),
        ("Review Period:", "January 2027 - December 2028", "Line Manager:", "David Osei"),
        ("Reviewed:","February 2027", "HR Partner:", "Sophia Reyes"),
    ]
    for r, (l1, v1, l2, v2) in enumerate(meta):
        row = table.rows[r]
        row.cells[0].text = l1
        row.cells[1].text = v1
        row.cells[2].text = l2
        row.cells[3].text = v2
    doc.add_paragraph()  # spacer

    _heading(doc, "1. Progress Review from 2025/2026 Development Plan", 1)
    _body(doc,
        "The following achievements against the 2025/2026 PDP goals are noted:\n\n"
        "COMPLETED GOALS:\n\n"
        "✓ Advanced Analytics (Goal 1)\n"
        "  Status: EXCEEDED\n"
        "  - Completed Advanced SQL course (Q1 2026) - distinction grade\n"
        "  - Delivered management reporting dashboard used by 4 BA practice teams\n"
        "  - Applied skills independently on three project engagements; data validation quality improved significantly\n"
        "  Assessment: Promoted to Proficient+ level in Advanced SQL. Ready for Lead BA data-analysis responsibilities.\n\n"
        "✓ CBAP Certification (Goal 2)\n"
        "  Status: COMPLETED\n"
        "  - Sat exam February 2026: PASSED (347/450 - well above pass mark of 280)\n"
        "  - CBAP credential now active since March 2026\n"
        "  - Applied BABOK knowledge in Meridian engagement (Phase 1 & 2)\n"
        "  Assessment: Formal senior BA credential achieved. Supports promotion readiness.\n\n"
        "✓ Solution Design Authorship (Goal 3)\n"
        "  Status: EXCEEDED\n"
        "  - Completed BCS International Diploma (Q3 2026) - Distinction\n"
        "  - Independently authored four solution design documents for 2026-2027 engagements\n"
        "  - Peer and manager feedback on solution design quality: Consistently strong, approved with minimal revision\n"
        "  Assessment: Now operates at Lead BA level for solution design. Primary promotion readiness gap closed.\n\n"
        "✓ Facilitation Confidence (Goal 4)\n"
        "  Status: EXCEEDED\n"
        "  - Completed 'Effective Facilitation' internal course (Q1 2026)\n"
        "  - Led 8 full facilitation sessions with director/C-suite stakeholders; peer feedback scores 4.8/5 average\n"
        "  - Meridian board presentation (Phase 2) was highlight engagement moment\n"
        "  Assessment: Stakeholder confidence in Emma's facilitation is now very high. Senior role-readiness demonstrated.\n\n"
        "✓ TOGAF 10 Foundation (Goal 5)\n"
        "  Status: COMPLETED\n"
        "  - Completed TOGAF 10 Foundation (Q1 2026) and passed exam - Certificate issued March 2026\n"
        "  - Attended Architecture Forum; developing enterprise architecture awareness\n"
        "  Assessment: Foundational architecture knowledge acquired. Positioned for future BA/Architect crossover roles.\n\n"
        "OVERALL 2025/2026 RESULT: All five goals achieved or exceeded. Progress demonstrates strong commitment to "
        "professional development and clear progression toward Lead BA / Senior Analyst role."
    )

    _heading(doc, "2. Current Skills Assessment - 2027", 1)
    _body(doc,
        "Updated skills assessment based on 2026 engagements and 2027 Meridian Phase 2 cycle:\n\n"
        "TECHNICAL SKILLS (Updated):\n"
        "  ✓ Requirements elicitation and documentation: ADVANCED / EXPERT\n"
        "  ✓ Business process modelling: ADVANCED\n"
        "  ✓ Solution design and architecture: ADVANCED (improved from Competent)\n"
        "  ✓ Agile / Scrum delivery: ADVANCED\n"
        "  ✓ Advanced SQL and data analysis: PROFICIENT+ (improved from Competent)\n"
        "  ✓ Power BI reporting and dashboards: PROFICIENT+ (improved from Competent)\n"
        "  ✓ UAT design and execution: ADVANCED\n"
        "  ✓ Stakeholder management and negotiation: PROFICIENT (strength area)\n"
        "  ✓ Workshop facilitation and design: ADVANCED (improved from Competent)\n\n"
        "DOMAIN KNOWLEDGE:\n"
        "  • Retail and e-commerce: ADVANCED (deepened on Meridian Phase 1 & 2)\n"
        "  • Supply chain operations: INTERMEDIATE (gained during Meridian WMS immersion)\n"
        "  • Financial services: ADVANCED (from three prior FS engagements)\n"
        "  • Logistics and inventory management: INTERMEDIATE\n\n"
        "SOFT SKILLS & LEADERSHIP:\n"
        "  • Written communication: STRONG / EXPERT\n"
        "  • Executive presentation skills: ADVANCED / NEAR-EXPERT\n"
        "  • Stakeholder leadership: ADVANCED\n"
        "  • Team mentoring: PROFICIENT+ (mentored 2 junior BAs in 2026-2027; feedback excellent)\n"
        "  • Cross-team collaboration: STRONG / EXPERT\n"
        "  • Change management: DEVELOPING (emerging strength; noted in Meridian feedback)"
    )

    _heading(doc, "3. Promotion Readiness Assessment", 1)
    _body(doc,
        "ASSESSMENT DATE: February 2027\n\n"
        "Emma is assessed as READY FOR PROMOTION TO LEAD BUSINESS ANALYST.\n\n"
        "Promotion Criteria Met:\n"
        "  ✓ Consistently delivers requirements at high quality with minimal rework\n"
        "  ✓ Manages complex stakeholder environments independently\n"
        "  ✓ Architects solutions for medium-to-large engagements without senior guidance\n"
        "  ✓ Mentors junior staff (demonstrated impact on two hires)\n"
        "  ✓ Contributes to BA discipline / community initiatives\n"
        "  ✓ Professional certifications in place: CBAP, BCS Diploma, TOGAF 10, IIBA membership\n"
        "  ✓ Client satisfaction / account feedback: Exceeds Expectations (Meridian, current reference)\n"
        "  ✓ Internal peer feedback: Strong across all competency areas\n"
        "  ✓ Manager recommendation: Strong promotion recommendation (David Osei, line manager)\n\n"
        "Promotion Timing: Recommended for promotion effective April 2027 (post-Q1 review cycle).\n"
        "Salary Adjustment: [HR to complete]\n"
        "Lead BA baseline: Senior + 15%\n"
        "Emma's case: Senior + 18% (due to strong delivery track record and growth trajectory)"
    )

    _heading(doc, "4. New Development Goals - 2027/2028", 1)
    _body(doc,
        "STRATEGIC CONTEXT:\n"
        "Following promotion to Lead BA, Emma's development focus shifts toward account leadership, "
        "delivery management, and strategic business acumen. She will take on mentoring responsibilities "
        "for up to 3 junior BAs.\n\n\n"
        "GOAL 1: Lead BA Account Delivery Excellence (New) [STRATEGIC]\n"
        "  Objective: Establish track record as lead BA on at least one stand-alone major engagement (>£500k)\n"
        "  Success Criteria:\n"
        "    • Meridian Phase 3 delivery: Lead BA role, delivered on-time/quality/budget\n"
        "    • Engagement margin target: >=28% (vs. 22% practice average for comparable scope)\n"
        "    • Client satisfaction: 4.5/5 or above\n"
        "    • Zero escalations for requirements/scope issues\n"
        "  Timeline: 2027-2028 (Meridian Phase 3, May 2027 - Dec 2027)\n"
        "  Measures: Engagement financial performance + client NPS + internal delivery metrics\n\n"
        "GOAL 2: Business Analyst Discipline Leadership [STRATEGIC]\n"
        "  Objective: Expand leadership of BA community; establish Emma as subject matter expert\n"
        "  Success Criteria:\n"
        "    • Chair BA Community of Practice (2028); establish monthly cadence\n"
        "    • Create new BA capability toolkit - requirements template library (20+ templates)\n"
        "    • Deliver 4 internal BA knowledge-share sessions (e.g., solution design, stakeholder mgmt) - Q2-Q4 2027\n"
        "    • Establish buddying group for 3 junior BAs (formal mentoring programme)\n"
        "  Timeline: 2027-2028, 5 h/month allocated\n"
        "  Measures: Community engagement + toolkit adoption + mentee feedback (360 review)\n\n"
        "GOAL 3: Senior Stakeholder & Account Leadership Skills [CAREER]\n"
        "  Objective: Develop comfort and competency working directly with C-suite / board-level stakeholders\n"
        "  Success Criteria:\n"
        "    • Lead at least one C-level investment decision workshop\n"
        "    • Complete senior leadership / executive communication course\n"
        "    • Deliver account strategy presentation (20 execs+) at least once\n"
        "    • Take on account planning responsibilities for one existing major client (Meridian likely)\n"
        "  Timeline: 2027-2028 (staggered across year)\n"
        "  Measures: Course completion + facilitation feedback + stakeholder feedback\n"
        "  Investment: Internal executive communication course (3 days, £2.5k)\n\n"
        "GOAL 4: Advanced Delivery & Programme Management Knowledge [SKILL]\n"
        "  Objective: Broaden understanding of delivery and programme interdependencies\n"
        "  Success Criteria:\n"
        "    • Complete PMI-PgMP Foundation or equivalent programme management awareness training\n"
        "    • Partner with engagement lead on one multi-workstream delivery; understand workstream dependencies\n"
        "    • Contribute to engagement resourcing and planning for 3 engagements (lead BA input)\n"
        "  Timeline: 2027-2028\n"
        "  Measures: Training completion + delivery planning contributions + PM feedback\n"
        "  Investment: PMI-PgMP Foundation course (2 days, £1.8k)\n\n"
        "GOAL 5: Industry Thought Leadership (Optional / Challenging) [CAREER]\n"
        "  Objective: Establish Emma as external thought leader on retail / supply chain BA practices\n"
        "  Success Criteria:\n"
        "    • Publish one article on retail business analysis for industry publication or conference proceedings\n"
        "    • Present at IIBA conference or industry event (BA-focused session)\n"
        "    • Contribute case study on Meridian engagement to published playbook (post-NDA clearance)\n"
        "  Timeline: 2027-2028 (stretch goal; not required for performance)\n"
        "  Measures: Article/conference acceptance + internal visibility + capability strengthening"
    )

    _heading(doc, "5. Development Activities & Timeline", 1)
    _body(doc,
        "FORMAL TRAINING & CERTIFICATIONS (2027-2028):\n"
        "  • Senior Executive Communication course (3 days) - Q2 2027 - £2.5k\n"
        "  • PMI-PgMP Foundation (2 days) - Q3 2027 - £1.8k\n"
        "  • IIBA Advanced Business Analysis certification track planning - Q4 2027\n"
        "  • Industry conference attendance (IIBA Europe Summit) - June 2028 - £1.2k + travel\n\n"
        "ON-THE-JOB LEARNING:\n"
        "  • Meridian Phase 3: Lead BA - requirements, stakeholder management, solution design\n"
        "  • Account strategy & planning: Support engagement lead on Meridian account plan (2 h/month)\n"
        "  • Multi-workstream delivery: Partner with PM on resource planning and workstream coordination\n"
        "  • Mentoring: Formal mentorship of 3 junior BAs (3 h/month structured, plus informal support)\n"
        "  • BA Community of Practice: Chair + 4 knowledge-share sessions (5 h/month + prep)\n\n"
        "EXTERNAL & COMMUNITY ACTIVITIES:\n"
        "  • IIBA UK Chapter: Maintain active membership; attend networking events (4 h/quarter)\n"
        "  • Thought leadership: Author case study article (Q3-Q4 2027, 20 h)\n"
        "  • Conference presentation proposal: Submit for IIBA 2028 conference (Q4 2027 deadline)\n\n"
        "TOTAL DEVELOPMENT INVESTMENT (2027-2028):\n"
        "  Training/courses: £5.5k\n"
        "  Conference/events: £1.2k\n"
        "  Time allocation: ~6-8 h/week (including mentoring, community, and on-the-job coaching)\n"
        "  Total: £6.7k + 312-416 hours (includes work time)"
    )

    _heading(doc, "6. Mentoring & Coaching Plan", 1)
    _body(doc,
        "MENTORING RESPONSIBILITIES (NEW - Post-promotion):\n\n"
        "Mentee 1: David Chen (Junior BA, ~2 years experience)\n"
        "  Current Focus: Solution design skills; independent engagement leadership\n"
        "  Mentoring Plan: Monthly 1-hr structured sessions; focus on design thinking and architecture\n"
        "  Expected Outcome: Ready for mid-level BA responsibilities by Q4 2027\n\n"
        "Mentee 2: [TBD - Q2 2027 assignent, likely new hire or second junior]\n"
        "  Current Focus: TBD post-identification\n"
        "  Mentoring Plan: Monthly 1-hr structured sessions; tailored to mentee development needs\n\n"
        "Mentee 3: [TBD - Q3 2027 assignent]\n"
        "  Current Focus: TBD post-identification\n"
        "  Mentoring Plan: Monthly 1-hr structured sessions; tailored to mentee development needs\n\n"
        "COACHING / LINE MANAGER SUPPORT:\n"
        "  Emma's line manager (David Osei) will provide:\n"
        "    • Quarterly leadership coaching (1 hour per quarter) focused on executive presence and stakeholder influence\n"
        "    • Sponsorship for stretch opportunities (large engagements, high-visibility projects)\n"
        "    • Feedback and accountability on development goals (mid-year, year-end reviews)\n"
        "  Coaching Investment: Internal (no external coach required at this stage)"
    )

    _heading(doc, "7. Performance Indicators & Review Checkpoints", 1)
    _body(doc,
        "MID-YEAR CHECKPOINT (July 2027):\n"
        "  • Meridian Phase 3 delivery progress: On track financially and schedule-wise?\n"
        "  • Executive communication course: Completed? Feedback positive?\n"
        "  • Account planning contributions: Visible on Meridian account?\n"
        "  • Mentee progress: On track with David Chen coaching?\n"
        "  • BA Community: Monthly cadence established?\n\n"
        "YEAR-END REVIEW (December 2027):\n"
        "  • Lead BA delivery: Phase 3 complete. Quality metrics, financial performance, client satisfaction?\n"
        "  • Lead BA readiness: Demonstrating competency at new level?\n"
        "  • Leadership contributions: Account leadership, community contributions, mentoring outcomes?\n"
        "  • Development progress: Certifications, courses, skill advancement?\n"
        "  • Career trajectory: On track for Principal/Senior Lead role by 2029?\n\n"
        "END OF PLAN (December 2028):\n"
        "  • Promotion impact: Has Emma successfully transitioned to Lead BA role?\n"
        "  • Account / discipline leadership: Track record on major engagements and community leadership?\n"
        "  • Next career steps: Principal BA readiness, Delivery Management, Account Lead transition, or other path?\n"
        "  • 2029 Development Plan: Informed by 2027-2028 performance and career intent"
    )

    _heading(doc, "8. Employee & Manager Sign-Off", 1)
    _body(doc,
        "This development plan has been jointly agreed between Emma Laurent, line manager David Osei, and HR Partner Sophia Reyes.\n\n"
        "Employee Signature: Emma Laurent                    Date: February 20, 2027\n"
        "  Emma confirms understanding of development expectations and commitment to 2027-2028 goals.\n\n"
        "Line Manager Signature: David Osei                  Date: February 22, 2027\n"
        "  David confirms endorsement of this plan and line manager support for development activities.\n\n"
        "HR Partner Signature: Sophia Reyes                  Date: February 24, 2027\n"
        "  HR confirms this plan aligns with promotion to Lead BA and organizational development priorities.\n\n"
        "Next Review: Mid-year checkpoint, July 2027"
    )

    doc.save(path)
    print(f"[OK] {path}")


# ========================================
# Main
# ========================================

if __name__ == "__main__":
    import os

    out_dir = "test_docs"
    os.makedirs(out_dir, exist_ok=True)

    print("Generating 2027 test documents for Emma Laurent...\n")
    
    create_client_feedback_2027(f"{out_dir}/Emma_Laurent_Client_Feedback_2027.docx")
    create_pdp_2027_2028(f"{out_dir}/Emma_Laurent_PDP_2027_2028.docx")

    print(f"\n✓ All 2027 test documents written to ./{out_dir}/")
    print("\nTo test the year linking feature:")
    print("  1. Upload Emma_Laurent_Client_Feedback_2027.docx with year=2027")
    print("  2. Upload Emma_Laurent_PDP_2027_2028.docx with year=2027")
    print("  3. View profile for Emma Laurent / 2027")
    print("  4. Check that year navigation shows: 2025 → 2027")
    print("  5. Verify previous_year_profile_id links 2027 to 2025 in database")
