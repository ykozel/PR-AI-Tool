"""
Script to generate test .docx files for the PR Profile pipeline.

Persona: Emma Laurent, Senior Business Analyst, Strategy & Innovation dept.
Documents cover two review years (2025 and 2026) so the comparison engine
can be tested end-to-end.

Run:
    pip install python-docx
    python generate_test_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------
# Helpers
# ---------------------------------------------

def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)
    return p


def _body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def _rating_question(
    doc: Document,
    number: int,
    title: str,
    options: list[str],
    selected: int,
    comment: str = "",
):
    """
    Render one rating feedback question.

    Args:
        number:   Question number (1-based)
        title:    Question label (may end with *)
        options:  Ordered list of rating options (first is always 'Can't respond')
        selected: 0-based index of the chosen option
        comment:  Optional free-text comment
    """
    # Question heading
    q_para = doc.add_paragraph()
    q_run = q_para.add_run(f"{number}. {title}")
    q_run.bold = True
    q_run.font.size = Pt(12)
    q_run.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)

    for idx, option in enumerate(options):
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.left_indent = Pt(18)
        if idx == selected:
            # Checked marker + bold text
            marker = p.add_run("[X] ")   # [X]
            marker.bold = True
            marker.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)
            chosen = p.add_run(option)
            chosen.bold = True
            chosen.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)
        else:
            marker = p.add_run("[ ] ")   # [ ]
            marker.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            rest = p.add_run(option)
            rest.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    # Comments field
    c_para = doc.add_paragraph()
    c_label = c_para.add_run("If you have additional comments, write them here:  ")
    c_label.italic = True
    c_label.font.size = Pt(10)
    c_label.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    if comment:
        c_val = c_para.add_run(comment)
        c_val.font.size = Pt(10)
    doc.add_paragraph()  # spacer


def _free_text_question(doc: Document, number: int, title: str, answer: str):
    """
    Render a free-text question with its filled-in answer.
    """
    q_para = doc.add_paragraph()
    q_run = q_para.add_run(f"{number}. {title}")
    q_run.bold = True
    q_run.font.size = Pt(12)
    q_run.font.color.rgb = RGBColor(0x1F, 0x47, 0x88)

    # Answer box
    a_para = doc.add_paragraph(style="List Paragraph")
    a_para.paragraph_format.left_indent = Pt(18)
    a_run = a_para.add_run(answer)
    a_run.font.size = Pt(11)
    doc.add_paragraph()  # spacer


def _meta_table(doc: Document, name: str, dept: str, role: str, period: str):
    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    labels = [
        ("Employee Name:", name, "Period:", period),
        ("Department:", dept, "Role:", role),
        ("Review Date:", "March 2026" if "2026" in period else "March 2025", "", ""),
    ]
    for row_idx, (l1, v1, l2, v2) in enumerate(labels):
        row = table.rows[row_idx]
        row.cells[0].text = l1
        row.cells[1].text = v1
        row.cells[2].text = l2
        row.cells[3].text = v2
    doc.add_paragraph()  # spacer


# ---------------------------------------------
# 2026 Documents  (current-year review)
# ---------------------------------------------
# 2025 Documents  (previous-year review, for comparison testing)
# ---------------------------------------------


# 
# Shared rating option lists
# 

_DELIVERY_APPROACH_OPTIONS = [
    "Architect-level: defines delivery approach and methodology across the organisation",
    "Advanced: independently designs project-level delivery approach",
    "Proficient: applies established methodology; seeks guidance on complex decisions",
    "Developing: follows prescribed approach with frequent guidance",
    "Foundational: requires close support to follow delivery process",
]

_TRANSPARENCY_OPTIONS = [
    "Exemplary: real-time dashboards and proactive escalation",
    "Strong: regular, accurate reporting with minimal prompting",
    "Adequate: shares status when asked; occasional gaps proactively",
    "Inconsistent: status updates require prompting; metrics unclear",
    "Insufficient: lack of transparency creates project risk",
]

_PROCESS_IMPROVEMENT_OPTIONS = [
    "Pioneer: drives organisation-wide process improvement initiatives",
    "Leader: regularly proposes and implements team-level improvements",
    "Contributor: identifies and raises improvements reactively",
    "Passive: rarely suggests improvements",
    "Not yet contributing to process improvement",
]

_TECH_SKILLS_OPTIONS = [
    "Expert: mentor to peers; evaluates and champions new tools",
    "Proficient: confident with current toolset; adapts to new tools independently",
    "Competent: effective with familiar tools; needs support for new ones",
    "Developing: requires regular guidance with tools",
    "Foundational: limited tool proficiency; needs significant support",
]

def create_company_function_feedback_2025(path: str):
    """Company function feedback for Emma Laurent - 2025 (single-question template)."""
    doc = Document()
    doc.add_heading("COMPANY FUNCTION FEEDBACK FORM - 2025", 0)
    _meta_table(
        doc,
        name="Emma Laurent",
        dept="Strategy & Innovation",
        role="Business Analyst",
        period="January - December 2025",
    )

    _free_text_question(
        doc, 1,
        "Please describe the employee's participation in the internal function "
        "in the recent period*",
        "Emma co-led the relaunch of the BA Community of Practice, which had been dormant "
        "for 18 months. She established a regular bi-monthly cadence, recruited 8 new members, "
        "and ran 4 sessions in 2025 with an average of 22 attendees. She created a shared "
        "resource library containing 15 BA and analysis templates now used by practitioners across the team.\n\n"
        "Emma also took on a buddy-mentor role in the Graduate Induction Programme, "
        "supporting two 2025 graduate joiners through their 6-month induction. Both graduates "
        "are now operating independently ahead of the programme's expectations.\n\n"
        "She presented the Meridian Retail engagement as a case study at the internal "
        "Project Showcase event, attended by 40 colleagues, and produced a lessons-learned "
        "document that has since been referenced in two subsequent client engagements."
    )

    doc.save(path)
    print(f"[OK]  {path}")


def create_auto_feedback_2025(path: str):
    """Auto (self) feedback form for Emma Laurent - 2025."""
    doc = Document()
    doc.add_heading("AUTO FEEDBACK FORM - ANNUAL REVIEW 2025", 0)
    _meta_table(
        doc,
        name="Emma Laurent",
        dept="Strategy & Innovation",
        role="Business Analyst",
        period="January - December 2025",
    )

    _heading(doc, "1. Previous Year Review Link", 1)
    _body(doc, (
        "Reference: 2024 Annual Review - Employee ID: EMP-2024-007832\n"
        "Previous Rating: Meets Expectations (3.6 / 5)\n"
        "Goals set in 2024:\n"
        "  - Achieve PMI-ACP certification (Achieved Oct 2025)\n"
        "  - Increase client-facing work to at least 50 % of billable time (Achieved - 62 %)\n"
        "  - Complete at least one cross-functional internal initiative (Achieved - BA CoP relaunch)"
    ))

    _heading(doc, "2. Certifications", 1)
    _body(doc, (
        "PMI Agile Certified Practitioner (PMI-ACP) - obtained October 2025\n"
        "Business Analysis Foundation (BCS) - obtained January 2025\n"
        "Microsoft Power BI Desktop - completed online certification June 2025\n\n"
        "In Progress:\n"
        "CBAP (Certified Business Analysis Professional) - exam booked February 2026\n"
        "TOGAF 10 Foundation - studying, exam planned Q1 2026"
    ))

    _heading(doc, "3. Learning & Development", 1)
    _body(doc, (
        "Courses Completed:\n"
        "  - Agile Business Analysis (20 h, Feb 2025)\n"
        "  - Data Visualisation with Power BI (18 h, May 2025)\n"
        "  - Effective Stakeholder Engagement (12 h, Jul 2025)\n"
        "  - Introduction to Enterprise Architecture (10 h, Sep 2025)\n\n"
        "Events:\n"
        "  - IIBA UK Chapter Annual Conference, London (Nov 2025)\n\n"
        "Total development hours in 2025: 120 hours\n\n"
        "Key Outcomes:\n"
        "  - Built Power BI dashboard used in Meridian engagement\n"
        "  - Applied agile BA techniques in Meridian project sprints"
    ))

    _heading(doc, "4. Self-Assessment & Feedback", 1)
    _body(doc, (
        "2025 was a year of consolidation and growth. The Meridian engagement pushed me out "
        "of my comfort zone into a more client-facing role and I responded positively. "
        "Obtaining the PMI-ACP was a significant milestone and has already improved the way "
        "I work in agile delivery teams.\n\n"
        "Proud of:\n"
        "  - Relaunching the BA Community of Practice (dormant for 18 months)\n"
        "  - Delivering the Meridian BRD that the client called their 'best ever'\n"
        "  - Supporting two junior analysts who both achieved their first professional certs\n\n"
        "Growth areas I am actively addressing:\n"
        "  - Facilitation confidence at exec level (practising in internal forums)\n"
        "  - Strategic business acumen (reading, courses planned)\n\n"
        "Self-rating: Exceeds Expectations (4.1 / 5)"
    ))

    _heading(doc, "5. Project Activity", 1)
    _body(doc, (
        "Meridian Retail Holdings - Omni-Channel Loyalty Platform (Mar-Nov 2025)\n"
        "Role: Business Analyst\n"
        "Team size: 12 people\n"
        "Key deliverables: Business Requirements Document, process maps, acceptance test scripts\n"
        "Outcome: Platform launched on target; enrolment conversion rate +18 %.\n\n"
        "Internal Data Quality Initiative (Jan-Mar 2025)\n"
        "Role: Analyst\n"
        "Conducted data-quality audit across four internal systems; produced remediation "
        "roadmap adopted by IT leadership.\n\n"
        "BA Community of Practice Relaunch (Apr-Dec 2025)\n"
        "Role: Co-Lead Organiser\n"
        "Relaunched dormant CoP, established regular cadence, recruited 8 new members."
    ))

    _heading(doc, "6. Function Activity", 1)
    _body(doc, (
        "BA Community of Practice - Co-Lead Organiser\n"
        "Relaunched CoP dormant since mid-2023. Ran 4 sessions in 2025 (avg. 22 attendees). "
        "Created shared resource library with 15 templates.\n\n"
        "Graduate Induction Programme - Buddy Mentor\n"
        "Mentored two 2025 graduate joiners through 6-month induction. Both now operating "
        "independently ahead of expectations.\n\n"
        "Involvement Level: Active Contributor / Emerging Leader"
    ))

    doc.save(path)
    print(f"[OK]  {path}")


def create_project_feedback_2025(path: str):
    """Project feedback form for Emma Laurent - 2025 (18-question template)."""
    doc = Document()
    doc.add_heading("PROJECT FEEDBACK FORM - 2025", 0)
    _meta_table(
        doc,
        name="Emma Laurent",
        dept="Strategy & Innovation",
        role="Business Analyst",
        period="January - December 2025",
    )

    _free_text_question(doc, 1, "What do you like about the person?*",
        "Emma brings solid documentation skills and genuine enthusiasm to the team. "
        "Her analysis deliverables are thorough and well-structured, and she builds good relationships "
        "with developers and product owners. She is always willing to support colleagues "
        "and has a can-do attitude even under pressure."
    )
    _free_text_question(doc, 2, "What could be improved?*",
        "Facilitation confidence in contested or executive-level meetings - Emma occasionally "
        "defers when she should lead. Delivery pace in the early sprints also had room for "
        "improvement, though this recovered strongly in H2."
    )
    _free_text_question(doc, 3,
        "Self Management: Being able to plan your activities & meet made commitments.*",
        "Generally good at planning and meeting commitments. A small number of early-sprint "
        "deliverables were submitted late; by Q3 Emma had resolved this and consistently "
        "met all deadlines. Continues to refine her estimation skills."
    )
    _free_text_question(doc, 4,
        "Analytical Skills/Product Mindset: Being able to see the bigger picture, understand the "
        "dependencies and take this into account when making the decisions. Gather small pieces "
        "into a single puzzle.*",
        "Good analytical skills with a growing product mindset. Occasionally focuses on "
        "detail before establishing the broader context - an area she is actively developing. "
        "The Power BI dashboard she produced gave the team a clearer picture of project progress."
    )
    _free_text_question(doc, 5,
        "Understanding of Software Development Lifecycle: Knowledge & understanding of "
        "development/testing methodologies.*",
        "Solid understanding of agile practices including sprint ceremonies and backlog "
        "refinement. Less experience with formal requirements frameworks and enterprise "
        "delivery methodologies - an acknowledged development area for 2026."
    )
    _free_text_question(doc, 6,
        "Flexibility/Agility: Ability to adapt to the situation. Balance, speed.*",
        "Adapts reasonably well to change. Required some support mid-project when scope "
        "shifted, but recovered well once the situation was clear and produced revised "
        "deliverable plans within the agreed timeframe."
    )
    _free_text_question(doc, 7,
        "Productivity/Velocity: How quickly the work is being done, consideration of the results "
        "achieved vs requirements & commitment.*",
        "Consistent, quality output once settled into the project rhythm. Velocity increased "
        "markedly in H2. The requirements documents she produced were praised by the client's engagement lead "
        "for their clarity and traceability."
    )
    _free_text_question(doc, 8,
        "Reliability/Commitment: Cares about the work to be done, being responsible for the "
        "commitment made.*",
        "Reliable and conscientious. Keeps commitments, escalates blockers promptly, and "
        "takes ownership of her deliverables. No commitments broken in H2."
    )
    _free_text_question(doc, 9,
        "Teamwork/Work Ethic: Being a team player, sharing knowledge, avoid blaming culture, "
        "focus on team results rather than just personal ones.*",
        "A genuine team player who actively supports colleagues and contributes to a "
        "positive team culture. Shared analysis templates that the team adopted across "
        "the project. Never observed deflecting blame."
    )
    _free_text_question(doc, 10,
        "Communication: Ability to listen & hear, ability to convey thoughts clearly.*",
        "Clear written communicator - analysis reports and status updates are well-structured "
        "and unambiguous. Verbal communication in larger meetings is an area for growth; "
        "Emma is actively working to build confidence in those settings."
    )
    _free_text_question(doc, 11, "English skills: Language skills*",
        "Proficient - strong written and spoken English with no communication barriers "
        "observed throughout the project."
    )
    _free_text_question(doc, 12,
        "Self Motivation & Control: Professional attitude towards work, where problems with "
        "motivation shall not take effect on performance & team's morale (self-control).*",
        "Consistently motivated and professional. Maintains a positive attitude and "
        "proactively looks for ways to add value, even during less interesting phases "
        "of the project."
    )
    _free_text_question(doc, 13,
        "Self Development/Professional Expertise: Constructive attitude towards criticism, "
        "looking for feedback. Learning new things to fill the gaps both in the "
        "day-to-day operation.*",
        "Receptive to feedback and acts on it. Obtained PMI-ACP this year and has CBAP "
        "booked for early 2026. Brings learning back to the team promptly - "
        "the Power BI dashboard is a direct example."
    )
    _free_text_question(doc, 14,
        "Any additional comments/feedback, please, give down below*",
        "Emma is growing into a strong mid-level business analyst. With continued development "
        "in solution design and facilitation confidence, she has excellent potential. "
        "We would welcome her on future engagements."
    )

    # Rating questions - 2025: solid mid-level performer, progressing but not yet senior
    _rating_question(doc, 15, "Delivery approach and project methodology*",
        _DELIVERY_APPROACH_OPTIONS, selected=2,
        comment="Emma has a solid grasp of delivery processes but benefited from senior guidance "
                "when defining the overall project approach. This is her clear focus for 2026."
    )
    _rating_question(doc, 16, "Transparency of project status and results*",
        _TRANSPARENCY_OPTIONS, selected=2,
        comment="Daily status was generally clear; occasional ad-hoc questions arose about "
                "progress metrics. The Power BI dashboard implemented in Q3 significantly "
                "improved transparency in H2."
    )
    _rating_question(doc, 17, "Process assessment and improvement*",
        _PROCESS_IMPROVEMENT_OPTIONS, selected=2,
        comment="Raised process improvements reactively when issues surfaced. Encouraged to "
                "establish a regular retrospective cadence to move towards proactive improvement."
    )
    _rating_question(doc, 18, "Technical skills*",
        _TECH_SKILLS_OPTIONS, selected=1,
        comment="Good foundational technical skills. Required support when evaluating new "
                "analytical tooling; grew noticeably more independent in this area by year-end."
    )

    doc.save(path)
    print(f"[OK]  {path}")


def create_pdp_2025(path: str):
    """Personal Development Plan for Emma Laurent - 2025/2026 cycle."""
    doc = Document()
    doc.add_heading("PERSONAL DEVELOPMENT PLAN - 2025 / 2026", 0)

    # Meta table
    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    meta = [
        ("Employee Name:", "Emma Laurent", "Employee ID:", "EMP-2024-007832"),
        ("Department:",    "Strategy & Innovation", "Role:",   "Senior Business Analyst"),
        ("Review Period:", "January 2025 - December 2026", "Line Manager:", "David Osei"),
    ]
    for r, (l1, v1, l2, v2) in enumerate(meta):
        row = table.rows[r]
        row.cells[0].text = l1
        row.cells[1].text = v1
        row.cells[2].text = l2
        row.cells[3].text = v2
    doc.add_paragraph()  # spacer

    _heading(doc, "1. Current Skills Assessment", 1)
    _body(doc, (
        "The following skills have been assessed as part of the 2025 annual review "
        "and form the baseline for this development plan.\n\n"
        "Technical Skills (confirmed proficient):\n"
        "  - Requirements elicitation and documentation: Advanced\n"
        "  - Business process modelling: Advanced\n"
        "  - Stakeholder analysis and management: Proficient\n"
        "  - Agile / Scrum delivery: Proficient\n"
        "  - SQL querying and data validation: Competent\n"
        "  - Power BI reporting and dashboards: Competent\n"
        "  - Workshop facilitation and design: Competent\n\n"
        "Domain Knowledge:\n"
        "  - Retail and e-commerce business analysis\n"
        "  - UAT co-ordination and sign-off management\n"
        "  - Requirements elicitation and BRD authoring\n\n"
        "Soft Skills:\n"
        "  - Written communication: Strong\n"
        "  - Stakeholder collaboration: Proficient\n"
        "  - Facilitation and presentation: Developing (target area)"
    ))

    _heading(doc, "2. Technical Skill Development Goals", 1)
    _body(doc, (
        "Goal 2a - Advanced Data Analysis and Reporting\n"
        "Target: Reach Proficient level in advanced SQL and Power BI by Q2 2026.\n"
        "Rationale: Data analysis capability is increasingly expected at Senior BA level. "
        "Stronger analytical skills will enable independent data-driven decision support "
        "and reduce dependency on data engineers for routine analysis.\n"
        "Actions:\n"
        "  - Complete 'Advanced SQL for Data Analysis' LinkedIn Learning path (15 h) - Q1 2026\n"
        "  - Build a management reporting dashboard for the BA practice\n"
        "  - Pair with the data team for two project workstreams\n\n"
        "Goal 2b - Advanced SQL and Data Analysis\n"
        "Target: Move from Competent to Proficient in SQL by Q3 2026.\n"
        "Rationale: Data validation gaps were identified on the Meridian project. "
        "Stronger SQL skills will allow independent back-end verification.\n"
        "Actions:\n"
        "  - Complete 'Advanced SQL for Data Analysis' LinkedIn Learning path (15 h) - Q2 2026\n"
        "  - Apply skills on at least two upcoming engagement data-validation workstreams\n\n"
        "Goal 2c - Solution Design and Requirements Architecture\n"
        "Target: Independently author project-level solution design documents by end of 2026.\n"
        "Rationale: Feedback in 2025 review highlighted that Emma benefited from senior guidance "
        "for solution scoping. This is the primary technical skill gap for promotion readiness.\n"
        "Actions:\n"
        "  - Study TOGAF and enterprise architecture fundamentals alongside CBAP exam prep\n"
        "  - Shadow a senior BA on two solution design workshops in Q1 2026\n"
        "  - Draft a solution design document independently for the next suitable engagement"
    ))

    _heading(doc, "3. Certifications and Qualifications", 1)
    _body(doc, (
        "Planned Certifications (2026):\n\n"
        "CBAP - Certified Business Analysis Professional (IIBA)\n"
        "  Status: Exam booked - February 2026\n"
        "  Study plan: 80 h self-study using BABOK Guide v3 and practice exams\n"
        "  Expected outcome: Validates senior BA capability; supports promotion case\n\n"
        "BCS International Diploma in Business Analysis\n"
        "  Status: Planned - exam target Q3 2026\n"
        "  Study plan: 60 h, BCS syllabus and practice exercises\n"
        "  Expected outcome: Closes solution design skill gap; positions Emma for senior and lead roles\n\n"
        "TOGAF 10 Foundation\n"
        "  Status: In progress - exam planned Q1 2026\n"
        "  Study plan: 30 h, official TOGAF courseware\n"
        "  Expected outcome: Broadens enterprise architecture awareness for BA and architecture crossover roles"
    ))

    _heading(doc, "4. Learning and Training Plan", 1)
    _body(doc, (
        "Formal Courses Planned:\n"
        "  - Advanced SQL for Data Analysis (LinkedIn Learning, 15 h) - Q1 2026\n"
        "  - BCS International Diploma in Business Analysis prep (BCS Learning, 24 h) - Q2 2026\n"
        "  - TOGAF 10 Foundation courseware (30 h, self-study) - Q1 2026\n"
        "  - Effective Facilitation and Presentation Skills (internal, 8 h) - Q1 2026\n"
        "  - Leadership Foundations for Senior ICs (internal, 12 h) - Q3 2026\n\n"
        "Conferences and Events:\n"
        "  - Business Analysis Summit Europe - June 2026 (approved)\n"
        "  - IIBA UK Chapter Annual Conference - November 2026 (planned)\n\n"
        "On-the-Job Learning:\n"
        "  - Senior BA shadowing (solution design workshops, two engagements) - H1 2026\n"
        "  - Management reporting dashboard build alongside data team - Q1 2026\n"
        "  - Present BA Community of Practice session on requirements techniques - Q2 2026\n\n"
        "Target development hours for 2026: 140 hours"
    ))

    _heading(doc, "5. Soft Skills and Leadership Development", 1)
    _body(doc, (
        "Goal 5a - Executive Facilitation Confidence\n"
        "Current level: Developing. Emma defers in contested senior meetings.\n"
        "Target: Comfortably facilitate workshops with director-level stakeholders by Q4 2026.\n"
        "Actions:\n"
        "  - Attend 'Effective Facilitation' internal course (Q1 2026)\n"
        "  - Volunteer to co-facilitate at least two internal steering group meetings\n"
        "  - Seek feedback after each facilitation session\n\n"
        "Goal 5b - Strategic Thinking and Business Acumen\n"
        "Current level: Competent at operational level; limited strategic framing.\n"
        "Target: Demonstrate commercial awareness in proposals and solution frameworks by Q3 2026.\n"
        "Actions:\n"
        "  - Read 'The Trusted Advisor' and 'Good Strategy / Bad Strategy'\n"
        "  - Apply commercial-impact framing in next two client-facing deliverables\n\n"
        "Goal 5c - Internal Mentoring and Knowledge Sharing\n"
        "Current level: Active - mentored two graduates in 2025.\n"
        "Target: Formalise mentoring role; take on one mid-level mentee in 2026.\n"
        "Actions:\n"
        "  - Register with internal mentoring matching programme by Q1 2026\n"
        "  - Run two BA CoP sessions focused on solution design and requirements techniques"
    ))

    _heading(doc, "6. Development Goals Summary (SMART)", 1)
    _body(doc, (
        "Goal 1: Advanced Analytics\n"
        "  Specific: Complete SQL training and deliver management reporting dashboard\n"
        "  Measurable: Dashboard used by at least two project teams\n"
        "  Achievable: 15 h course + 20 h build within Q1 spare capacity\n"
        "  Relevant: Core capability gap for promotion to Lead BA\n"
        "  Time-bound: Complete by end of Q1 2026\n\n"
        "Goal 2: CBAP Certification\n"
        "  Specific: Pass CBAP exam on first attempt\n"
        "  Measurable: Pass mark achieved; certificate issued\n"
        "  Achievable: 80 h study plan; exam booked\n"
        "  Relevant: Validates senior BA competency\n"
        "  Time-bound: Exam date February 2026\n\n"
        "Goal 3: Solution Design Authorship\n"
        "  Specific: Independently draft a full solution design document for a real engagement\n"
        "  Measurable: Document approved by lead BA without major structural changes\n"
        "  Achievable: Following shadow programme and TOGAF/BCS study\n"
        "  Relevant: Closes primary promotion-readiness gap\n"
        "  Time-bound: First draft submitted by Q3 2026\n\n"
        "Goal 4: Facilitation Confidence\n"
        "  Specific: Lead at least two director-level facilitation sessions without co-facilitator\n"
        "  Measurable: Peer feedback score >= 4/5 on facilitation effectiveness\n"
        "  Achievable: Post internal course and practice opportunities\n"
        "  Relevant: Required for senior stakeholder management\n"
        "  Time-bound: Complete by Q4 2026"
    ))

    _heading(doc, "7. Sign-Off", 1)
    _body(doc, (
        "Employee: Emma Laurent                  Date: January 15, 2026\n"
        "Line Manager: David Osei                Date: January 20, 2026\n"
        "HR Partner: Sophia Reyes                Date: January 22, 2026\n\n"
        "Review checkpoint scheduled: June 2026 (mid-year check-in)"
    ))

    doc.save(path)
    print(f"[OK]  {path}")


# ---------------------------------------------
# Main
# ---------------------------------------------

if __name__ == "__main__":
    import os

    out_dir = "test_docs"
    os.makedirs(out_dir, exist_ok=True)

    print("Generating 2025 documents...")
    create_company_function_feedback_2025(f"{out_dir}/Emma_Laurent_Company_Function_Feedback_2025.docx")
    create_auto_feedback_2025(f"{out_dir}/Emma_Laurent_Auto_Feedback_2025.docx")
    create_project_feedback_2025(f"{out_dir}/Emma_Laurent_Project_Feedback_2025.docx")
    create_pdp_2025(f"{out_dir}/Emma_Laurent_PDP_2025_2026.docx")

    print(f"\n  All 4 test documents written to ./{out_dir}/")
